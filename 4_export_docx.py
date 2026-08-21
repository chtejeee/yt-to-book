"""Step 4 — Combine chapters into a single DOCX book."""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE = Path(__file__).parent
VIDEOS_PATH = BASE / "data" / "videos.json"
CHAPTERS_DIR = BASE / "data" / "chapters"
IMAGES_DIR = BASE / "data" / "images"
OUT_PATH = BASE / "output" / "book.docx"


def parse_chapter(text: str) -> tuple[str, str]:
    lines = text.strip().splitlines()
    if lines and lines[0].startswith("#"):
        title = lines[0].lstrip("#").strip()
        body = "\n".join(lines[1:]).strip()
    else:
        title = lines[0].strip() if lines else "Untitled Chapter"
        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
    return title, body


def load_images(video_id: str) -> tuple[float, list[dict]]:
    manifest = IMAGES_DIR / video_id / "images.json"
    if not manifest.exists():
        return 0.0, []
    data = json.loads(manifest.read_text())
    return data.get("duration", 0.0), data.get("images", [])


def compute_placements(num_blocks: int, images: list[dict], duration: float) -> dict[int, list[dict]]:
    """Map each image to a paragraph-block index, by matching its proportional
    position in the video's timeline to the same proportional position in the
    chapter text. Approximate — the chapter is a rewritten summary, not a
    transcript — but keeps images near roughly the right content."""
    placements: dict[int, list[dict]] = {}
    if not images or duration <= 0 or num_blocks == 0:
        return placements
    for img in images:
        proportion = min(max(img["timestamp"] / duration, 0.0), 1.0)
        idx = min(int(round(proportion * num_blocks)), num_blocks - 1)
        placements.setdefault(idx, []).append(img)
    return placements


def add_markdown_paragraph(doc: Document, text: str):
    """Add a paragraph, converting **bold** / *italic* markers into real run formatting."""
    p = doc.add_paragraph()
    for token in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            p.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            p.add_run(token[1:-1]).italic = True
        else:
            p.add_run(token)
    return p


def add_image(doc: Document, video_id: str, img: dict, figure_no: int):
    path = IMAGES_DIR / video_id / img["file"]
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(4.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figure {figure_no}. {img['caption']}")
    run.italic = True
    run.font.size = Pt(9.5)


def add_body_paragraphs(doc: Document, body: str, video_id: str, placements: dict, fig_counter: list[int]):
    blocks = [b.strip() for b in body.split("\n\n") if b.strip()]
    for i, block in enumerate(blocks):
        for img in placements.get(i, []):
            add_image(doc, video_id, img, fig_counter[0])
            fig_counter[0] += 1

        heading_match = re.match(r"^#{2,3}\s+(.*)", block)
        # Local models tend to emit "**Subheading**" rather than "## Subheading" — treat a
        # block that's a single bold-wrapped line the same way as a markdown heading.
        bold_heading_match = re.match(r"^\*\*(.+?)\*\*$", block) if not heading_match else None
        if heading_match:
            doc.add_heading(heading_match.group(1).strip(), level=2)
        elif bold_heading_match:
            doc.add_heading(bold_heading_match.group(1).strip(), level=2)
        else:
            add_markdown_paragraph(doc, block)

    for img in placements.get(len(blocks), []):
        add_image(doc, video_id, img, fig_counter[0])
        fig_counter[0] += 1


def main():
    if not VIDEOS_PATH.exists():
        raise SystemExit("data/videos.json not found — run 1_fetch_videos.py first")

    videos = {v["id"]: v for v in json.loads(VIDEOS_PATH.read_text())}
    chapter_files = sorted(CHAPTERS_DIR.glob("*.txt"))

    if not chapter_files:
        raise SystemExit("No chapters found in data/chapters/ — run 3_build_book.py first")

    def sort_key(path: Path):
        v = videos.get(path.stem, {})
        return v.get("upload_date") or "00000000"

    chapter_files.sort(key=sort_key)

    doc = Document()

    # Cover page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("The Book")
    run.font.size = Pt(36)
    run.bold = True
    doc.add_page_break()

    # Table of contents (chapter titles only — generated after we parse chapters)
    toc_heading = doc.add_heading("Table of Contents", level=1)
    doc.add_page_break()

    chapters = []
    for path in chapter_files:
        video_id = path.stem
        text = path.read_text(encoding="utf-8")
        title, body = parse_chapter(text)
        chapters.append((video_id, title, body))

    fig_counter = [1]
    for video_id, title, body in chapters:
        doc.add_heading(title, level=1)
        duration, images = load_images(video_id)
        num_blocks = len([b for b in body.split("\n\n") if b.strip()])
        placements = compute_placements(num_blocks, images, duration)
        add_body_paragraphs(doc, body, video_id, placements, fig_counter)
        doc.add_page_break()

    # Appendix of video URLs
    doc.add_heading("Appendix: Source Videos", level=1)
    for video_id, title, _ in chapters:
        url = videos.get(video_id, {}).get("url", f"https://www.youtube.com/watch?v={video_id}")
        p = doc.add_paragraph()
        p.add_run(f"{title}").bold = True
        doc.add_paragraph(url)

    # Fill in TOC now that we know chapter titles
    toc_paragraphs = [f"{i+1}. {title}" for i, (_, title, _) in enumerate(chapters)]
    for line in reversed(toc_paragraphs):
        toc_heading._p.addnext(doc.add_paragraph(line)._p)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Book exported -> {OUT_PATH}  ({len(chapters)} chapters, {fig_counter[0] - 1} figures)")


if __name__ == "__main__":
    main()
